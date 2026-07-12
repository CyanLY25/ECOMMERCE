import pandas as pd
import numpy as np
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List

from ia.config.config import AIConfig
from ia.utils.logger import setup_logger
from ia.reports.report_data import load_report_data
from ia.reports import interpretations as interp

logger = setup_logger("report_generator")

# Título y descripción usados en portada / metadatos de los 3 formatos
REPORT_TITLE = "Predicción de Demanda E-Commerce"
REPORT_SUBTITLE = "Reporte Técnico: EDA, Modelos de Redes Neuronales, Validación Cruzada, Tuning y Pruebas Estadísticas"


class ReportGenerator:
    """
    Clase para generar reportes de los resultados de los modelos de ML.
    Soporta formatos PDF, Excel y Word.
    """

    def __init__(self, config: AIConfig):
        """
        Inicializa el generador de reportes con la configuración.
        
        Args:
            config: Objeto de configuración.
        """
        self.config = config

    def generate_html_eda_report(self, df: pd.DataFrame, statistics: pd.DataFrame, outlier_stats: Dict[str, Any]) -> None:
        """
        Genera un reporte HTML de EDA.
        
        Args:
            df: DataFrame procesado.
            statistics: DataFrame con estadísticas descriptivas.
            outlier_stats: Diccionario con estadísticas de outliers.
        """
        logger.info("Generando reporte HTML de EDA...")
        
        figures_dir = self.config.FIGURES_DIR
        figure_files = sorted(figures_dir.glob("*.png"))
        
        # Generar HTML
        html_content = f"""
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <title>Reporte de Análisis Exploratorio de Datos</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #3498db; margin-top: 40px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .figure {{ margin: 30px 0; }}
                .figure img {{ max-width: 100%; height: auto; }}
            </style>
        </head>
        <body>
            <h1>Reporte de Análisis Exploratorio de Datos (EDA)</h1>
            
            <h2>Resumen del Dataset</h2>
            <p><strong>Número de registros:</strong> {len(df)}</p>
            <p><strong>Número de columnas:</strong> {len(df.columns)}</p>
            <p><strong>Columnas:</strong> {', '.join(df.columns)}</p>
            
            <h2>Valores Nulos</h2>
            <table>
                <tr><th>Columna</th><th>Valores Nulos</th><th>Porcentaje</th></tr>
                {''.join([f'<tr><td>{col}</td><td>{df[col].isna().sum()}</td><td>{(df[col].isna().sum()/len(df)*100):.2f}%</td></tr>' for col in df.columns])}
            </table>
            
            <h2>Estadísticas Descriptivas</h2>
            {statistics.to_html(index=False, classes='table')}
            
            <h2>Outliers</h2>
            <table>
                <tr><th>Columna</th><th>Cantidad de Outliers</th><th>Porcentaje</th></tr>
                {''.join([f'<tr><td>{col}</td><td>{stats["count_outliers"]}</td><td>{stats["percentage_outliers"]:.2f}%</td></tr>' for col, stats in outlier_stats.items()])}
            </table>
            
            <h2>Visualizaciones</h2>
            {''.join([f'<div class="figure"><h3>{fig.stem}</h3><img src="figures/{fig.name}" alt="{fig.stem}"></div>' for fig in figure_files])}
            
            <h2>Conclusiones Automáticas</h2>
            <ul>
                <li>El dataset contiene {len(df)} registros válidos después de la limpieza.</li>
                <li>Se detectaron outliers en las columnas: {', '.join(outlier_stats.keys())}.</li>
                <li>Las variables numéricas principales son: {', '.join(df.select_dtypes(include=[np.number]).columns.tolist())}.</li>
            </ul>
            
        </body>
        </html>
        """
        
        # Guardar HTML
        report_path = self.config.EDA_REPORT_PATH
        report_path.write_text(html_content, encoding="utf-8")
        logger.info(f"Reporte HTML guardado en {report_path}")

    def generate_pdf(self, data: Optional[Dict[str, Any]] = None, output_path: Optional[Path] = None) -> Path:
        """
        Genera el reporte final en formato PDF: EDA, comparación de modelos,
        validación cruzada, hiperparámetros y pruebas estadísticas, con
        tablas, figuras y su interpretación.

        Args:
            data: Diccionario devuelto por `load_report_data`. Si es None, se carga
                automáticamente desde las rutas configuradas en AIConfig.
            output_path: Ruta donde guardar el PDF. Si es None, usa config.REPORT_PDF_PATH.

        Returns:
            La ruta del PDF generado.
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            Image as RLImage, PageBreak,
        )

        logger.info("Generando reporte PDF...")
        data = data or load_report_data(self.config)
        output_path = Path(output_path) if output_path else self.config.REPORT_PDF_PATH
        output_path.parent.mkdir(parents=True, exist_ok=True)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="H1c", parent=styles["Heading1"], textColor=colors.HexColor("#2c3e50")))
        styles.add(ParagraphStyle(name="H2c", parent=styles["Heading2"], textColor=colors.HexColor("#3498db"), spaceBefore=16))
        styles.add(ParagraphStyle(name="Interp", parent=styles["BodyText"], textColor=colors.HexColor("#333333"),
                                   backColor=colors.HexColor("#f4f6f7"), borderPadding=6, spaceBefore=6, spaceAfter=12))

        story: List[Any] = []

        # ---------- Portada ----------
        story.append(Spacer(1, 4 * cm))
        story.append(Paragraph(REPORT_TITLE, styles["Title"]))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(REPORT_SUBTITLE, styles["Heading3"]))
        story.append(Spacer(1, 1 * cm))
        story.append(Paragraph(f"Fecha de generación: {data.get('generated_at', datetime.now())}", styles["Normal"]))
        story.append(PageBreak())

        def df_to_table(df: pd.DataFrame, max_rows: int = 25, float_fmt: str = "{:.3f}") -> Table:
            df = df.copy().head(max_rows)
            for col in df.select_dtypes(include=[np.number]).columns:
                df[col] = df[col].map(lambda v: float_fmt.format(v) if pd.notnull(v) else "")
            table_data = [list(df.columns)] + df.astype(str).values.tolist()
            t = Table(table_data, repeatRows=1, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            return t

        def add_section(title: str, table_obj: Optional[Table], interpretation: str, note_if_missing: str):
            story.append(Paragraph(title, styles["H2c"]))
            if table_obj is not None:
                story.append(table_obj)
            else:
                story.append(Paragraph(note_if_missing, styles["Italic"]))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretation}", styles["Interp"]))

        def add_figure(fig_path: Path, caption: str):
            try:
                img = RLImage(str(fig_path), width=15 * cm, height=9 * cm, kind="proportional")
                story.append(img)
                if caption:
                    story.append(Paragraph(f"<i>{caption}</i>", styles["Normal"]))
                story.append(Spacer(1, 0.4 * cm))
            except Exception as e:
                logger.warning(f"No se pudo incrustar la figura {fig_path}: {e}")

        # ---------- 1. EDA ----------
        story.append(Paragraph("1. Análisis Exploratorio de Datos (EDA)", styles["H1c"]))
        eda_table = df_to_table(data["eda_statistics"]) if data.get("eda_statistics") is not None else None
        add_section(
            "1.1 Estadísticos Descriptivos", eda_table,
            interp.interpret_eda(data.get("eda_statistics")),
            "No se encontró statistics.csv."
        )
        for fig_name in ["5_correlation_heatmap.png", "13_time_series_sales.png", "11_country_distribution.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        story.append(PageBreak())

        # ---------- 2. Entrenamiento y comparación de modelos ----------
        story.append(Paragraph("2. Entrenamiento y Comparación de Modelos", styles["H1c"]))
        story.append(Paragraph(
            "Se entrenaron tres modelos clásicos (MLP, LSTM, GRU) y dos modelos híbridos "
            "(CNN-LSTM, CNN-GRU), guardando el mejor modelo en formato .keras para su consumo "
            "directo desde el backend sin necesidad de reentrenamiento.", styles["Normal"]
        ))
        cmp_table = df_to_table(data["model_comparison"]) if data.get("model_comparison") is not None else None
        add_section(
            "2.1 Tabla Comparativa (MAE, RMSE, MSE, MAPE, R², Tiempo, Épocas)", cmp_table,
            interp.interpret_model_comparison(data.get("model_comparison"), data.get("best_model")),
            "No se encontró model_comparison.csv."
        )
        for fig_name in ["cross_validation_boxplot.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        story.append(PageBreak())

        # ---------- 3. Validación Cruzada ----------
        story.append(Paragraph(f"3. Validación Cruzada ({self.config.CV_FOLDS} folds)", styles["H1c"]))
        cv_table = df_to_table(data["cv_summary"]) if data.get("cv_summary") is not None else None
        add_section(
            "3.1 Resumen por Modelo (media ± desviación estándar)", cv_table,
            interp.interpret_cross_validation(data.get("cv_summary")),
            "No se encontró cross_validation_summary.csv."
        )
        for fig_name in ["cross_validation_rmse.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        story.append(PageBreak())

        # ---------- 4. Hiperparámetros ----------
        story.append(Paragraph("4. Ajuste de Hiperparámetros (Tuning)", styles["H1c"]))
        story.append(Paragraph(f"<b>Interpretación:</b> {interp.interpret_hyperparameters(data.get('best_hyperparameters'))}".replace("\n", "<br/>"), styles["Interp"]))
        for fig_name in ["tuning_comparison.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        story.append(PageBreak())

        # ---------- 5. Pruebas Estadísticas ----------
        story.append(Paragraph("5. Pruebas Estadísticas Robustas", styles["H1c"]))
        rank_table = df_to_table(data["ranking"]) if data.get("ranking") is not None else None
        add_section(
            "5.1 Ranking de Modelos (Friedman / RMSE)", rank_table,
            interp.interpret_statistics(data.get("friedman"), data.get("wilcoxon"), data.get("nemenyi"), data.get("ranking")),
            "No se encontró ranking_results.csv."
        )
        wil_table = df_to_table(data["wilcoxon"], max_rows=10) if data.get("wilcoxon") is not None else None
        if wil_table is not None:
            story.append(Paragraph("5.2 Comparaciones Pareadas (Wilcoxon, corrección Bonferroni)", styles["H2c"]))
            story.append(wil_table)
        for fig_name in ["significance_heatmap.png", "critical_difference.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        story.append(PageBreak())

        # ---------- 6. Conclusiones ----------
        story.append(Paragraph("6. Conclusiones", styles["H1c"]))
        conclusiones = data.get("statistical_conclusions") or "No se encontró statistical_conclusions.md."
        for line in conclusiones.split("\n"):
            if line.strip():
                story.append(Paragraph(line, styles["Normal"]))

        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                 topMargin=2 * cm, bottomMargin=2 * cm,
                                 leftMargin=2 * cm, rightMargin=2 * cm,
                                 title=REPORT_TITLE)
        doc.build(story)
        logger.info(f"Reporte PDF guardado en {output_path}")
        return output_path

    def generate_excel(self, data: Optional[Dict[str, Any]] = None, output_path: Optional[Path] = None) -> Path:
        """
        Genera el reporte final en formato Excel (.xlsx), con una hoja por
        sección (Resumen, EDA, Comparación de Modelos, Cross Validation,
        Hiperparámetros, Pruebas Estadísticas), cada tabla acompañada de su
        interpretación y las figuras clave incrustadas en la hoja de Resumen.

        Args:
            data: Diccionario devuelto por `load_report_data`. Si es None, se carga
                automáticamente desde las rutas configuradas en AIConfig.
            output_path: Ruta donde guardar el .xlsx. Si es None, usa config.REPORT_EXCEL_PATH.

        Returns:
            La ruta del .xlsx generado.
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
        from openpyxl.drawing.image import Image as XLImage

        logger.info("Generando reporte Excel...")
        data = data or load_report_data(self.config)
        output_path = Path(output_path) if output_path else self.config.REPORT_EXCEL_PATH
        output_path.parent.mkdir(parents=True, exist_ok=True)

        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        wrap = Alignment(wrap_text=True, vertical="top")

        wb = Workbook()

        def write_df(ws, df: pd.DataFrame, start_row: int = 1) -> int:
            for j, col in enumerate(df.columns, start=1):
                cell = ws.cell(row=start_row, column=j, value=str(col))
                cell.font = header_font
                cell.fill = header_fill
            for i, (_, row) in enumerate(df.iterrows(), start=start_row + 1):
                for j, val in enumerate(row, start=1):
                    if isinstance(val, (np.floating, float)):
                        val = round(float(val), 4)
                    elif isinstance(val, (np.integer,)):
                        val = int(val)
                    ws.cell(row=i, column=j, value=val)
            for j, col in enumerate(df.columns, start=1):
                ws.column_dimensions[get_column_letter(j)].width = max(12, min(30, len(str(col)) + 4))
            return start_row + len(df) + 1

        def write_interpretation(ws, row: int, text: str, ncols: int = 6) -> int:
            cell = ws.cell(row=row, column=1, value=f"Interpretación: {text}")
            cell.alignment = wrap
            cell.font = Font(italic=True)
            ws.merge_cells(start_row=row, start_column=1, end_row=row + 3, end_column=max(ncols, 4))
            ws.row_dimensions[row].height = 90
            return row + 5

        # ---------- Hoja Resumen ----------
        ws = wb.active
        ws.title = "Resumen"
        ws["A1"] = REPORT_TITLE
        ws["A1"].font = Font(bold=True, size=16)
        ws["A2"] = REPORT_SUBTITLE
        ws["A2"].font = Font(italic=True, size=11, color="3498DB")
        ws["A3"] = f"Fecha de generación: {data.get('generated_at', datetime.now())}"
        ws["A4"] = f"CV folds: {self.config.CV_FOLDS}"
        best_model = data.get("best_model") or {}
        row = 6
        if best_model:
            ws.cell(row=row, column=1, value="Mejor modelo (RMSE):").font = Font(bold=True)
            ws.cell(row=row, column=2, value=best_model.get("model"))
            ws.cell(row=row + 1, column=1, value="RMSE")
            ws.cell(row=row + 1, column=2, value=round(best_model.get("rmse", 0), 4))
            ws.cell(row=row + 2, column=1, value="MAE")
            ws.cell(row=row + 2, column=2, value=round(best_model.get("mae", 0), 4))
            ws.cell(row=row + 3, column=1, value="R²")
            ws.cell(row=row + 3, column=2, value=round(best_model.get("r2", 0), 4))
            row += 5
        row = write_interpretation(ws, row, interp.interpret_model_comparison(data.get("model_comparison"), best_model))
        # Insertar figura principal en el resumen
        fig_path = data["figures"].get("5_correlation_heatmap.png")
        if fig_path:
            try:
                img = XLImage(str(fig_path))
                img.width, img.height = 480, 300
                ws.add_image(img, f"A{row}")
            except Exception as e:
                logger.warning(f"No se pudo incrustar imagen en Excel: {e}")

        # ---------- Hoja EDA ----------
        ws_eda = wb.create_sheet("EDA")
        if data.get("eda_statistics") is not None:
            r = write_df(ws_eda, data["eda_statistics"])
        else:
            ws_eda["A1"] = "No se encontró statistics.csv."
            r = 3
        write_interpretation(ws_eda, r, interp.interpret_eda(data.get("eda_statistics")))

        # ---------- Hoja Comparación de Modelos ----------
        ws_cmp = wb.create_sheet("Comparacion_Modelos")
        if data.get("model_comparison") is not None:
            r = write_df(ws_cmp, data["model_comparison"])
        else:
            ws_cmp["A1"] = "No se encontró model_comparison.csv."
            r = 3
        write_interpretation(ws_cmp, r, interp.interpret_model_comparison(data.get("model_comparison"), best_model))

        # ---------- Hoja Cross Validation ----------
        ws_cv = wb.create_sheet("Cross_Validation")
        if data.get("cv_summary") is not None:
            r = write_df(ws_cv, data["cv_summary"])
        else:
            ws_cv["A1"] = "No se encontró cross_validation_summary.csv."
            r = 3
        write_interpretation(ws_cv, r, interp.interpret_cross_validation(data.get("cv_summary")))

        # ---------- Hoja Hiperparámetros ----------
        ws_hp = wb.create_sheet("Hiperparametros")
        best_hp = data.get("best_hyperparameters") or {}
        r = 1
        for modelo, params in best_hp.items():
            ws_hp.cell(row=r, column=1, value=modelo).font = Font(bold=True)
            r += 1
            for k, v in params.items():
                ws_hp.cell(row=r, column=1, value=k)
                ws_hp.cell(row=r, column=2, value=v)
                r += 1
            r += 1
        write_interpretation(ws_hp, r, interp.interpret_hyperparameters(best_hp))

        # ---------- Hoja Pruebas Estadísticas ----------
        ws_stats = wb.create_sheet("Pruebas_Estadisticas")
        r = 1
        for label, key in [("Friedman", "friedman"), ("Ranking", "ranking"),
                            ("Nemenyi", "nemenyi"), ("Wilcoxon", "wilcoxon")]:
            df = data.get(key)
            ws_stats.cell(row=r, column=1, value=label).font = Font(bold=True, size=12)
            r += 1
            if df is not None and not df.empty:
                r = write_df(ws_stats, df, start_row=r)
            else:
                ws_stats.cell(row=r, column=1, value=f"No se encontró {key}_results.csv")
                r += 2
        write_interpretation(
            ws_stats, r,
            interp.interpret_statistics(data.get("friedman"), data.get("wilcoxon"), data.get("nemenyi"), data.get("ranking"))
        )

        wb.save(str(output_path))
        logger.info(f"Reporte Excel guardado en {output_path}")
        return output_path

    def generate_all(self, data: Optional[Dict[str, Any]] = None) -> Dict[str, Path]:
        """
        Genera los tres reportes finales (PDF, Word, Excel) en una sola
        llamada, reutilizando los mismos datos cargados una única vez.

        Returns:
            Diccionario {"pdf": ruta, "word": ruta, "excel": ruta}.
        """
        data = data or load_report_data(self.config)
        return {
            "pdf": self.generate_pdf(data),
            "word": self.generate_word(data),
            "excel": self.generate_excel(data),
        }

    def generate_word(self, data: Optional[Dict[str, Any]] = None, output_path: Optional[Path] = None) -> Path:
        """
        Genera el reporte final en formato Word (.docx), con la misma
        estructura que el PDF: EDA, comparación de modelos, validación
        cruzada, hiperparámetros, pruebas estadísticas y conclusiones,
        incluyendo tablas, figuras y su interpretación.

        Args:
            data: Diccionario devuelto por `load_report_data`. Si es None, se carga
                automáticamente desde las rutas configuradas en AIConfig.
            output_path: Ruta donde guardar el .docx. Si es None, usa config.REPORT_WORD_PATH.

        Returns:
            La ruta del .docx generado.
        """
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        logger.info("Generando reporte Word...")
        data = data or load_report_data(self.config)
        output_path = Path(output_path) if output_path else self.config.REPORT_WORD_PATH
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        def add_heading(text: str, level: int = 1):
            h = doc.add_heading(text, level=level)
            return h

        def add_interpretation(text: str):
            p = doc.add_paragraph()
            run = p.add_run("Interpretación: ")
            run.bold = True
            p.add_run(text)
            p.paragraph_format.space_after = Pt(12)

        def add_df_table(df: pd.DataFrame, max_rows: int = 25, float_fmt: str = "{:.3f}"):
            df = df.copy().head(max_rows)
            for col in df.select_dtypes(include=[np.number]).columns:
                df[col] = df[col].map(lambda v: float_fmt.format(v) if pd.notnull(v) else "")
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            return table

        def add_figure(fig_path: Path, caption: str):
            try:
                doc.add_picture(str(fig_path), width=Cm(15))
                if caption:
                    cap = doc.add_paragraph()
                    cap_run = cap.add_run(caption)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
            except Exception as e:
                logger.warning(f"No se pudo incrustar la figura {fig_path}: {e}")

        # ---------- Portada ----------
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(REPORT_TITLE)
        title_run.bold = True
        title_run.font.size = Pt(26)

        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_p.add_run(REPORT_SUBTITLE)
        subtitle_run.font.size = Pt(13)
        subtitle_run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Fecha de generación: {data.get('generated_at', datetime.now())}")
        doc.add_page_break()

        # ---------- 1. EDA ----------
        add_heading("1. Análisis Exploratorio de Datos (EDA)", 1)
        add_heading("1.1 Estadísticos Descriptivos", 2)
        if data.get("eda_statistics") is not None:
            add_df_table(data["eda_statistics"])
        else:
            doc.add_paragraph("No se encontró statistics.csv.")
        add_interpretation(interp.interpret_eda(data.get("eda_statistics")))
        for fig_name in ["5_correlation_heatmap.png", "13_time_series_sales.png", "11_country_distribution.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        doc.add_page_break()

        # ---------- 2. Comparación de modelos ----------
        add_heading("2. Entrenamiento y Comparación de Modelos", 1)
        doc.add_paragraph(
            "Se entrenaron tres modelos clásicos (MLP, LSTM, GRU) y dos modelos híbridos "
            "(CNN-LSTM, CNN-GRU), guardando el mejor modelo en formato .keras para su consumo "
            "directo desde el backend sin necesidad de reentrenamiento."
        )
        add_heading("2.1 Tabla Comparativa (MAE, RMSE, MSE, MAPE, R², Tiempo, Épocas)", 2)
        if data.get("model_comparison") is not None:
            add_df_table(data["model_comparison"])
        else:
            doc.add_paragraph("No se encontró model_comparison.csv.")
        add_interpretation(interp.interpret_model_comparison(data.get("model_comparison"), data.get("best_model")))
        fig_path = data["figures"].get("cross_validation_boxplot.png")
        if fig_path:
            add_figure(fig_path, interp.interpret_figure("cross_validation_boxplot.png"))
        doc.add_page_break()

        # ---------- 3. Validación Cruzada ----------
        add_heading(f"3. Validación Cruzada ({self.config.CV_FOLDS} folds)", 1)
        add_heading("3.1 Resumen por Modelo (media ± desviación estándar)", 2)
        if data.get("cv_summary") is not None:
            add_df_table(data["cv_summary"])
        else:
            doc.add_paragraph("No se encontró cross_validation_summary.csv.")
        add_interpretation(interp.interpret_cross_validation(data.get("cv_summary")))
        fig_path = data["figures"].get("cross_validation_rmse.png")
        if fig_path:
            add_figure(fig_path, interp.interpret_figure("cross_validation_rmse.png"))
        doc.add_page_break()

        # ---------- 4. Hiperparámetros ----------
        add_heading("4. Ajuste de Hiperparámetros (Tuning)", 1)
        add_interpretation(interp.interpret_hyperparameters(data.get("best_hyperparameters")))
        fig_path = data["figures"].get("tuning_comparison.png")
        if fig_path:
            add_figure(fig_path, interp.interpret_figure("tuning_comparison.png"))
        doc.add_page_break()

        # ---------- 5. Pruebas Estadísticas ----------
        add_heading("5. Pruebas Estadísticas Robustas", 1)
        add_heading("5.1 Ranking de Modelos (Friedman / RMSE)", 2)
        if data.get("ranking") is not None:
            add_df_table(data["ranking"])
        else:
            doc.add_paragraph("No se encontró ranking_results.csv.")
        add_interpretation(interp.interpret_statistics(data.get("friedman"), data.get("wilcoxon"), data.get("nemenyi"), data.get("ranking")))
        if data.get("wilcoxon") is not None:
            add_heading("5.2 Comparaciones Pareadas (Wilcoxon, corrección Bonferroni)", 2)
            add_df_table(data["wilcoxon"], max_rows=10)
        for fig_name in ["significance_heatmap.png", "critical_difference.png"]:
            fig_path = data["figures"].get(fig_name)
            if fig_path:
                add_figure(fig_path, interp.interpret_figure(fig_name))
        doc.add_page_break()

        # ---------- 6. Conclusiones ----------
        add_heading("6. Conclusiones", 1)
        conclusiones = data.get("statistical_conclusions") or "No se encontró statistical_conclusions.md."
        for line in conclusiones.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.save(str(output_path))
        logger.info(f"Reporte Word guardado en {output_path}")
        return output_path
